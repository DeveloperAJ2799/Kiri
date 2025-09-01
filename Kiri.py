# main.py
import sys
import os
import asyncio
import json
import datetime
from pathlib import Path
import fitz  # PyMuPDF
from ollama import AsyncClient
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout,
    QPushButton, QProgressBar, QTextEdit, QFileDialog, QTreeWidget, QTreeWidgetItem,
    QLabel, QMessageBox, QTabWidget, QComboBox, QSplitter, QHeaderView, QLineEdit, QSpinBox, QInputDialog
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QDir
from PyQt5.QtGui import QFont, QIcon, QColor, QPalette, QTextCursor # Import QTextCursor
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SummarizationWorker(QThread):
    progress = pyqtSignal(int)
    # Use a more general signal for output
    output = pyqtSignal(str)
    finished_signal = pyqtSignal(str)

    def __init__(self, text_data, model="gemma3:4b-it-qat", prompt_type="summary", custom_prompt="", answer_length="", is_chunked=False):
        super().__init__()
        self.text_data = text_data # Can be full text (str) or chunks (list)
        self.model = model
        self.prompt_type = prompt_type
        self.custom_prompt = custom_prompt
        self.answer_length = answer_length
        self.is_chunked = is_chunked # Flag to indicate if text_data is a list of chunks
        self.client = AsyncClient()
        self.result = "" # Store the final combined result

        # Different prompts for different content types
        self.prompts = {
            "summary": (
                "You are an expert note-taking assistant specializing in creating structured, "
                "detailed summaries. Create a comprehensive summary of the provided text. "
                "Focus on key facts, important concepts, and critical takeaways. "
                "Organize information logically with clear sections and paragraphs. "
                "Do not use markdown or special formatting in your response."
            ),
            "brief": (
                "Create a concise brief overview of the following text. "
                "Focus on the most important concepts and key takeaways. "
                "Use a clear paragraph structure."
            ),
            "qna": (
                "Generate 5 important questions and detailed answers based on the following text. "
                "Format each pair clearly as:\n"
                "Q: [Question]\n"
                "A: [Detailed answer]\n"
                "Separate each pair with a blank line."
            ),
            "questions": (
                "Create 10 practice questions based on the following text. "
                "Include a mix of multiple choice, short answer, and discussion questions. "
                "Number each question clearly."
            ),
            "professor": self.custom_prompt if self.custom_prompt else (
                "You are a helpful professor assistant. Answer questions about the provided text "
                "in a clear, educational manner. Provide detailed explanations and examples where appropriate."
            )
        }

    async def process_full_text(self, text):
        """Process the entire text as a single unit."""
        try:
            self.output.emit("Sending full text to model...")
            base_prompt = self.prompts.get(self.prompt_type, self.prompts["summary"])

            # Modify prompt based on answer length requirement (mainly for professor mode)
            if self.answer_length and self.prompt_type == "professor":
                length_prompt = f"Answer length: {self.answer_length}. "
                prompt = length_prompt + base_prompt
            else:
                prompt = base_prompt

            response = await self.client.chat(
                model=self.model,
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": text} # Send the whole text
                ],
                options={
                    "temperature": 0.4 if self.prompt_type in ["summary", "brief"] else 0.7,
                    "top_p": 0.9
                }
            )
            self.result = response['message']['content']
            self.progress.emit(100)
            self.output.emit("Processing complete.")
        except Exception as e:
            error_msg = f"[ERROR] Failed to process text: {str(e)}"
            self.result = error_msg
            self.output.emit(error_msg)

    async def process_chunks(self, chunks):
        """Process text in chunks (fallback or for specific content types)."""
        self.output.emit(f"Split into {len(chunks)} chunks for processing...")
        semaphore = asyncio.Semaphore(2) # Limit concurrent requests
        summaries = []

        async def process_chunk(chunk, index):
            async with semaphore:
                try:
                    self.output.emit(f"Processing chunk {index+1}/{len(chunks)}...")
                    base_prompt = self.prompts.get(self.prompt_type, self.prompts["summary"])

                    if self.answer_length and self.prompt_type == "professor":
                        length_prompt = f"Answer length: {self.answer_length}. "
                        prompt = length_prompt + base_prompt
                    else:
                        prompt = base_prompt

                    response = await self.client.chat(
                        model=self.model,
                        messages=[
                            {"role": "system", "content": prompt},
                            {"role": "user", "content": chunk}
                        ],
                        options={
                            "temperature": 0.4 if self.prompt_type in ["summary", "brief"] else 0.7,
                            "top_p": 0.9
                        }
                    )
                    summary = response['message']['content']
                    summaries.append(summary)
                    self.progress.emit(int((index + 1) / len(chunks) * 100))
                except Exception as e:
                    error_msg = f"[ERROR] Processing chunk {index+1}: {str(e)}"
                    self.output.emit(error_msg)
                    summaries.append(error_msg)

        tasks = [process_chunk(chunk, i) for i, chunk in enumerate(chunks)]
        await asyncio.gather(*tasks)

        # Combine results
        if self.prompt_type == "summary":
            # For summaries, try to create a cohesive whole (this is simplistic)
            self.result = "\n\n--- SECTION BREAK ---\n\n".join(summaries)
        else:
            # For other types, just join
            self.result = "\n\n".join(summaries)

    def run(self):
        try:
            if isinstance(self.text_data, str) and not self.is_chunked:
                # Run the single text processing coroutine
                asyncio.run(self.process_full_text(self.text_data))
            elif isinstance(self.text_data, list) or self.is_chunked:
                # Run the chunked processing coroutine
                chunks = self.text_data if isinstance(self.text_data, list) else [self.text_data]
                asyncio.run(self.process_chunks(chunks))
            else:
                raise ValueError("Invalid text_data type for processing")
        except Exception as e:
            self.result = f"[FATAL ERROR] {str(e)}"
        self.finished_signal.emit(self.result)


class PDFManagerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("kiri")
        self.setGeometry(100, 100, 1100, 800)
        # --- Black and Blue Theme ---
        self.setStyleSheet(self.get_black_blue_stylesheet())
        # Apply dark palette for better theme consistency
        self.set_dark_palette()
        # Data storage
        self.notes_file = "pdf_notes.json"
        self.notes = self.load_notes()
        self.current_pdf_text = ""
        # Store the PDF filename for export
        self.current_pdf_filename = ""
        # Store the generated topic for export
        self.current_pdf_topic = ""
        # Create central widget and layout
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)
        self.main_layout.setSpacing(10)
        self.main_layout.setContentsMargins(15, 15, 15, 15)
        # Header
        self.create_header()
        # Tab widget
        self.tab_widget = QTabWidget()
        self.main_layout.addWidget(self.tab_widget)
        # Create tabs
        self.create_file_browser_tab()
        self.create_processor_tab()
        self.create_notes_tab()
        self.create_content_generator_tab()
        self.create_professor_mode_tab()
        # Status bar
        self.statusBar().showMessage("Ready")

    def get_black_blue_stylesheet(self):
        return """
        /* Main window background */
        QMainWindow, QWidget {
            background-color: #000000; /* Black */
            color: #00ccff; /* Bright Blue text */
        }
        /* Buttons */
        QPushButton {
            background-color: #003366; /* Dark Blue */
            color: #ffffff; /* White text */
            border: 1px solid #0055aa; /* Medium Blue border */
            padding: 10px 15px;
            border-radius: 6px;
            font-weight: 600;
            font-size: 13px;
        }
        QPushButton:hover {
            background-color: #0055aa; /* Medium Blue */
        }
        QPushButton:pressed {
            background-color: #0077cc; /* Lighter Blue */
        }
        QPushButton:disabled {
            background-color: #333333; /* Dark Gray */
            color: #666666; /* Light Gray text */
            border: 1px solid #444444;
        }
        /* Text Edits */
        QTextEdit {
            background-color: #001122; /* Very Dark Blue-Black */
            color: #00ccff; /* Bright Blue text */
            border: 1px solid #004477; /* Darker Blue border */
            border-radius: 6px;
            padding: 10px;
            font-family: 'Segoe UI', sans-serif;
            font-size: 11pt;
            selection-background-color: #0055aa; /* Medium Blue */
            selection-color: #ffffff; /* White */
        }
        QTextEdit:disabled {
            background-color: #000a11; /* Even darker */
            color: #0088aa; /* Muted Blue */
        }
        /* Progress Bar */
        QProgressBar {
            border: 1px solid #004477;
            border-radius: 6px;
            text-align: center;
            height: 20px;
            background-color: #001122; /* Dark background */
        }
        QProgressBar::chunk {
            background-color: #0066cc; /* Blue chunk */
            border-radius: 5px;
        }
        /* Labels */
        QLabel {
            color: #00ccff; /* Bright Blue text */
            font-size: 13px;
        }
        /* Tab Widget */
        QTabWidget::pane {
            border: 1px solid #004477;
            border-radius: 6px;
            padding: 5px;
            background-color: #000000; /* Black */
        }
        QTabBar::tab {
            background: #002244; /* Very Dark Blue */
            border: 1px solid #004477;
            padding: 8px 15px;
            margin-right: 2px;
            border-top-left-radius: 4px;
            border-top-right-radius: 4px;
            color: #00ccff; /* Bright Blue text */
        }
        QTabBar::tab:selected {
            background: #003366; /* Dark Blue */
            color: #ffffff; /* White text */
            font-weight: bold;
        }
        QTabBar::tab:hover:!selected {
            background: #004477; /* Slightly lighter on hover */
        }
        /* Tree Widgets */
        QTreeWidget {
            background-color: #001122; /* Very Dark Blue-Black */
            alternate-background-color: #000a11; /* Slightly different shade */
            color: #00ccff; /* Bright Blue text */
            border: 1px solid #004477;
            border-radius: 6px;
            gridline-color: #003366; /* Dark Blue lines */
        }
        QTreeWidget::item {
            padding: 3px;
        }
        QTreeWidget::item:selected {
            background-color: #003366; /* Dark Blue */
            color: #ffffff; /* White text */
        }
        QTreeWidget::item:hover {
            background-color: #002244; /* Slightly lighter on hover */
        }
        QHeaderView::section {
            background-color: #002244; /* Very Dark Blue header */
            color: #00ccff; /* Bright Blue text */
            padding: 5px;
            border: 1px solid #003366;
            font-weight: bold;
        }
        /* Combo Boxes */
        QComboBox {
            background-color: #002244; /* Very Dark Blue */
            color: #00ccff; /* Bright Blue text */
            border: 1px solid #004477;
            border-radius: 4px;
            padding: 5px;
            min-height: 20px;
        }
        QComboBox:disabled {
            background-color: #001122;
            color: #0088aa;
        }
        QComboBox QAbstractItemView {
            background-color: #001122; /* Dropdown background */
            border: 1px solid #004477;
            selection-background-color: #003366; /* Dark Blue selection */
            selection-color: #ffffff; /* White text */
        }
        /* Spin Boxes */
        QSpinBox {
            background-color: #002244; /* Very Dark Blue */
            color: #00ccff; /* Bright Blue text */
            border: 1px solid #004477;
            border-radius: 4px;
            padding: 5px;
        }
        QSpinBox:disabled {
            background-color: #001122;
            color: #0088aa;
        }
        QSpinBox::up-button, QSpinBox::down-button {
            background-color: #003366; /* Dark Blue buttons */
            border: 1px solid #004477;
            subcontrol-origin: border;
            width: 16px;
        }
        QSpinBox::up-button:hover, QSpinBox::down-button:hover {
            background-color: #0055aa; /* Medium Blue on hover */
        }
        /* Status Bar */
        QStatusBar {
            background-color: #001122; /* Very Dark Blue-Black */
            color: #00ccff; /* Bright Blue text */
            border-top: 1px solid #003366;
        }
        /* Scrollbars */
        QScrollBar:vertical {
            background-color: #001122;
            width: 15px;
            margin: 15px 3px 15px 3px;
            border: 1px solid #003366;
            border-radius: 4px;
        }
        QScrollBar::handle:vertical {
            background-color: #003366;
            min-height: 20px;
            border-radius: 4px;
        }
        QScrollBar::handle:vertical:hover {
            background-color: #0055aa;
        }
        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
            background-color: #002244;
            height: 12px;
            subcontrol-origin: margin;
            border: 1px solid #003366;
        }
        QScrollBar:horizontal {
            background-color: #001122;
            height: 15px;
            margin: 3px 15px 3px 15px;
            border: 1px solid #003366;
            border-radius: 4px;
        }
        QScrollBar::handle:horizontal {
            background-color: #003366;
            min-width: 20px;
            border-radius: 4px;
        }
        QScrollBar::handle:horizontal:hover {
            background-color: #0055aa;
        }
        QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
            background-color: #002244;
            width: 12px;
            subcontrol-origin: margin;
            border: 1px solid #003366;
        }
        """

    def set_dark_palette(self):
        """Apply a dark palette for better theme consistency."""
        dark_palette = QPalette()
        dark_palette.setColor(QPalette.Window, QColor(0, 0, 0)) # Black
        dark_palette.setColor(QPalette.WindowText, QColor(0, 204, 255)) # Bright Blue
        dark_palette.setColor(QPalette.Base, QColor(0, 17, 34)) # Very Dark Blue-Black
        dark_palette.setColor(QPalette.AlternateBase, QColor(0, 10, 17)) # Slightly different shade
        dark_palette.setColor(QPalette.ToolTipBase, QColor(0, 204, 255)) # Bright Blue
        dark_palette.setColor(QPalette.ToolTipText, QColor(0, 0, 0)) # Black
        dark_palette.setColor(QPalette.Text, QColor(0, 204, 255)) # Bright Blue
        dark_palette.setColor(QPalette.Disabled, QPalette.Text, QColor(0, 136, 170)) # Muted Blue
        dark_palette.setColor(QPalette.Button, QColor(0, 51, 102)) # Dark Blue
        dark_palette.setColor(QPalette.ButtonText, QColor(255, 255, 255)) # White
        dark_palette.setColor(QPalette.BrightText, QColor(255, 0, 0)) # Red (for errors/alerts)
        dark_palette.setColor(QPalette.Link, QColor(0, 102, 204)) # Blue Link
        dark_palette.setColor(QPalette.Highlight, QColor(0, 85, 170)) # Medium Blue
        dark_palette.setColor(QPalette.HighlightedText, QColor(255, 255, 255)) # White
        dark_palette.setColor(QPalette.PlaceholderText, QColor(0, 136, 170)) # Muted Blue
        self.setPalette(dark_palette)

    def create_header(self):
        header_layout = QVBoxLayout()
        # Title
        header_label = QLabel("Kiri")
        header_font = QFont("Arial", 22, QFont.Bold)
        header_label.setFont(header_font)
        header_label.setStyleSheet("color: #00ccff; margin: 10px 0;") # Bright Blue
        header_label.setAlignment(Qt.AlignCenter)
        # Info
        info_label = QLabel("Made by Ayaan Jiwani")
        info_label.setAlignment(Qt.AlignCenter)
        info_label.setStyleSheet("color: #0088aa; font-size: 12px; margin-bottom: 15px;") # Muted Blue
        header_layout.addWidget(header_label)
        header_layout.addWidget(info_label)
        self.main_layout.addLayout(header_layout)

    def create_file_browser_tab(self):
        self.browser_tab = QWidget()
        self.tab_widget.addTab(self.browser_tab, "File Browser")
        layout = QVBoxLayout(self.browser_tab)
        # Directory selection
        dir_layout = QHBoxLayout()
        self.dir_button = QPushButton("Select Directory")
        self.dir_button.clicked.connect(self.select_directory)
        self.dir_label = QLabel("No directory selected")
        self.dir_label.setStyleSheet("color: #0088aa; font-style: italic;") # Muted Blue
        self.refresh_button = QPushButton("Refresh")
        self.refresh_button.clicked.connect(self.refresh_file_tree)
        dir_layout.addWidget(self.dir_button)
        dir_layout.addWidget(self.dir_label)
        dir_layout.addWidget(self.refresh_button)
        layout.addLayout(dir_layout)
        # File tree
        self.file_tree = QTreeWidget()
        self.file_tree.setHeaderLabels(["Name", "Type", "Size"])
        self.file_tree.header().setSectionResizeMode(0, QHeaderView.Stretch)
        self.file_tree.header().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.file_tree.header().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.file_tree.itemDoubleClicked.connect(self.on_item_double_clicked)
        layout.addWidget(self.file_tree)
        # Action buttons
        action_layout = QHBoxLayout()
        self.open_pdf_button = QPushButton("Open PDF")
        self.open_pdf_button.clicked.connect(self.open_selected_pdf)
        self.open_pdf_button.setEnabled(False)
        self.process_pdf_button = QPushButton("Process PDF")
        self.process_pdf_button.clicked.connect(self.process_selected_pdf)
        self.process_pdf_button.setEnabled(False)
        action_layout.addWidget(self.open_pdf_button)
        action_layout.addWidget(self.process_pdf_button)
        layout.addLayout(action_layout)

    def create_processor_tab(self):
        self.processor_tab = QWidget()
        self.tab_widget.addTab(self.processor_tab, "Process PDF")
        layout = QVBoxLayout(self.processor_tab)
        # File selection
        file_layout = QHBoxLayout()
        self.file_button = QPushButton("Select PDF File")
        self.file_button.clicked.connect(self.select_file)
        self.file_label = QLabel("No file selected")
        self.file_label.setStyleSheet("color: #0088aa; font-style: italic;") # Muted Blue
        file_layout.addWidget(self.file_button)
        file_layout.addWidget(self.file_label)
        layout.addLayout(file_layout)
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)
        # Process button
        self.process_button = QPushButton("Generate Summary with Kiri")
        self.process_button.clicked.connect(self.process_pdf)
        self.process_button.setEnabled(False)
        layout.addWidget(self.process_button)
        # Preview area
        preview_label = QLabel("Summary Preview:")
        preview_label.setStyleSheet("font-weight: 600; margin-top: 10px; color: #00ccff;") # Bright Blue
        self.preview_area = QTextEdit()
        self.preview_area.setReadOnly(True)
        layout.addWidget(preview_label)
        layout.addWidget(self.preview_area)
        # Export button
        self.export_button = QPushButton("Export Summary as Editable Word Document")
        self.export_button.clicked.connect(self.export_word)
        self.export_button.setEnabled(False)
        layout.addWidget(self.export_button)

    def create_notes_tab(self):
        self.notes_tab = QWidget()
        self.tab_widget.addTab(self.notes_tab, "Saved Notes")
        layout = QHBoxLayout(self.notes_tab)
        # Notes list
        notes_layout = QVBoxLayout()
        notes_label = QLabel("Saved Notes:")
        notes_label.setStyleSheet("font-weight: 600; color: #00ccff;") # Bright Blue
        notes_layout.addWidget(notes_label)
        self.notes_list = QTreeWidget()
        self.notes_list.setHeaderLabels(["Filename", "Date", "ID"])
        self.notes_list.header().setSectionResizeMode(0, QHeaderView.Stretch)
        self.notes_list.itemSelectionChanged.connect(self.on_note_selected)
        notes_layout.addWidget(self.notes_list)
        # Notes buttons
        notes_buttons_layout = QHBoxLayout()
        self.refresh_notes_button = QPushButton("Refresh")
        self.refresh_notes_button.clicked.connect(self.refresh_notes_list)
        self.delete_note_button = QPushButton("Delete")
        self.delete_note_button.clicked.connect(self.delete_note)
        self.delete_note_button.setEnabled(False)
        notes_buttons_layout.addWidget(self.refresh_notes_button)
        notes_buttons_layout.addWidget(self.delete_note_button)
        notes_layout.addLayout(notes_buttons_layout)
        # Note content
        content_layout = QVBoxLayout()
        content_label = QLabel("Note Content:")
        content_label.setStyleSheet("font-weight: 600; color: #00ccff;") # Bright Blue
        content_layout.addWidget(content_label)
        self.note_content = QTextEdit()
        self.note_content.setReadOnly(True)
        content_layout.addWidget(self.note_content)
        self.export_note_button = QPushButton("Export to Word")
        self.export_note_button.clicked.connect(self.export_note_word)
        self.export_note_button.setEnabled(False)
        content_layout.addWidget(self.export_note_button)
        layout.addLayout(notes_layout, 1)
        layout.addLayout(content_layout, 2)
        self.refresh_notes_list()

    def create_content_generator_tab(self):
        self.generator_tab = QWidget()
        self.tab_widget.addTab(self.generator_tab, "Content Generator")
        layout = QVBoxLayout(self.generator_tab)
        # File selection
        file_layout = QHBoxLayout()
        self.content_file_button = QPushButton("Select PDF File")
        self.content_file_button.clicked.connect(self.select_content_file)
        self.content_file_label = QLabel("No file selected")
        self.content_file_label.setStyleSheet("color: #0088aa; font-style: italic;") # Muted Blue
        file_layout.addWidget(self.content_file_button)
        file_layout.addWidget(self.content_file_label)
        layout.addLayout(file_layout)
        # Content type selection
        type_layout = QHBoxLayout()
        type_label = QLabel("Content Type:")
        type_label.setStyleSheet("color: #00ccff;") # Bright Blue
        type_layout.addWidget(type_label)
        self.content_type_combo = QComboBox()
        self.content_type_combo.addItems(["Summary", "Brief Overview", "Q&A", "Practice Questions"])
        self.content_type_combo.currentTextChanged.connect(self.on_content_type_changed)
        type_layout.addWidget(self.content_type_combo)
        self.generate_content_button = QPushButton("Generate Content")
        self.generate_content_button.clicked.connect(self.generate_content)
        self.generate_content_button.setEnabled(False)
        type_layout.addWidget(self.generate_content_button)
        layout.addLayout(type_layout)
        # Question parameters (only visible for question types)
        self.question_params_widget = QWidget()
        question_params_layout = QHBoxLayout(self.question_params_widget)
        self.question_count_label = QLabel("Number of Questions:")
        self.question_count_label.setStyleSheet("color: #00ccff;") # Bright Blue
        self.question_count_spin = QSpinBox()
        self.question_count_spin.setRange(1, 50)
        self.question_count_spin.setValue(5)
        self.question_type_label = QLabel("Question Type:")
        self.question_type_label.setStyleSheet("color: #00ccff;") # Bright Blue
        self.question_type_combo = QComboBox()
        self.question_type_combo.addItems(["Multiple Choice", "Short Answer", "Discussion", "Mixed"])
        question_params_layout.addWidget(self.question_count_label)
        question_params_layout.addWidget(self.question_count_spin)
        question_params_layout.addWidget(self.question_type_label)
        question_params_layout.addWidget(self.question_type_combo)
        self.question_params_widget.setVisible(False)
        layout.addWidget(self.question_params_widget)
        # Content preview
        preview_label = QLabel("Generated Content:")
        preview_label.setStyleSheet("font-weight: 600; margin-top: 10px; color: #00ccff;") # Bright Blue
        self.content_preview = QTextEdit()
        self.content_preview.setReadOnly(True)
        layout.addWidget(preview_label)
        layout.addWidget(self.content_preview)
        # Export button
        self.export_content_button = QPushButton("Export Content")
        self.export_content_button.clicked.connect(self.export_content)
        self.export_content_button.setEnabled(False)
        layout.addWidget(self.export_content_button)

    def create_professor_mode_tab(self):
        self.professor_tab = QWidget()
        self.tab_widget.addTab(self.professor_tab, "Professor Mode")
        layout = QVBoxLayout(self.professor_tab)
        # PDF selection for professor mode
        pdf_layout = QHBoxLayout()
        self.professor_pdf_button = QPushButton("Select PDF for Q&A")
        self.professor_pdf_button.clicked.connect(self.select_professor_pdf)
        self.professor_pdf_label = QLabel("No PDF selected")
        self.professor_pdf_label.setStyleSheet("color: #0088aa; font-style: italic;") # Muted Blue
        pdf_layout.addWidget(self.professor_pdf_button)
        pdf_layout.addWidget(self.professor_pdf_label)
        layout.addLayout(pdf_layout)
        # Answer length control
        length_layout = QHBoxLayout()
        self.answer_length_label = QLabel("Answer Length:")
        self.answer_length_label.setStyleSheet("color: #00ccff;") # Bright Blue
        self.answer_length_combo = QComboBox()
        self.answer_length_combo.addItems(["Brief (1-2 sentences)", "Medium (3-5 sentences)", "Detailed (1 paragraph)", "Comprehensive (2+ paragraphs)"])
        length_layout.addWidget(self.answer_length_label)
        length_layout.addWidget(self.answer_length_combo)
        layout.addLayout(length_layout)
        # Question input and answer section
        qa_layout = QVBoxLayout()
        question_label = QLabel("Ask a Question:")
        question_label.setStyleSheet("font-weight: 600; color: #00ccff;") # Bright Blue
        qa_layout.addWidget(question_label)
        self.question_input = QTextEdit()
        self.question_input.setMaximumHeight(60)
        self.question_input.setPlaceholderText("Type your question here...")
        qa_layout.addWidget(self.question_input)
        answer_buttons_layout = QHBoxLayout()
        self.ask_question_button = QPushButton("Ask Question")
        self.ask_question_button.clicked.connect(self.ask_question)
        self.ask_question_button.setEnabled(False)
        self.clear_qa_button = QPushButton("Clear")
        self.clear_qa_button.clicked.connect(self.clear_qa)
        answer_buttons_layout.addWidget(self.ask_question_button)
        answer_buttons_layout.addWidget(self.clear_qa_button)
        qa_layout.addLayout(answer_buttons_layout)
        layout.addLayout(qa_layout)
        # Answer area
        answer_label = QLabel("Answer:")
        answer_label.setStyleSheet("font-weight: 600; margin-top: 10px; color: #00ccff;") # Bright Blue
        self.answer_area = QTextEdit()
        self.answer_area.setReadOnly(True)
        layout.addWidget(answer_label)
        layout.addWidget(self.answer_area)

    def on_content_type_changed(self, text):
        # Show/hide question parameters based on content type
        is_question_type = text in ["Q&A", "Practice Questions"]
        self.question_params_widget.setVisible(is_question_type)

    def select_directory(self):
        directory = QFileDialog.getExistingDirectory(self, "Select Directory")
        if directory:
            self.dir_label.setText(directory)
            self.populate_file_tree(directory)

    def populate_file_tree(self, directory):
        self.file_tree.clear()
        self.populate_tree_item(self.file_tree.invisibleRootItem(), Path(directory))

    def populate_tree_item(self, parent, path):
        try:
            for item in path.iterdir():
                if item.is_dir():
                    dir_item = QTreeWidgetItem(parent, [item.name, "Directory", ""])
                    dir_item.setChildIndicatorPolicy(QTreeWidgetItem.ShowIndicator)
                    dir_item.setData(0, Qt.UserRole, str(item))
                elif item.suffix.lower() == '.pdf':
                    size = item.stat().st_size
                    size_str = self.format_file_size(size)
                    file_item = QTreeWidgetItem(parent, [item.name, "PDF", size_str])
                    file_item.setData(0, Qt.UserRole, str(item))
        except PermissionError:
            pass

    def format_file_size(self, size_bytes):
        if size_bytes == 0:
            return "0 B"
        size_names = ["B", "KB", "MB", "GB"]
        i = 0
        while size_bytes >= 1024 and i < len(size_names) - 1:
            size_bytes /= 1024.0
            i += 1
        return f"{size_bytes:.1f} {size_names[i]}"

    def refresh_file_tree(self):
        directory = self.dir_label.text()
        if directory and directory != "No directory selected":
            self.populate_file_tree(directory)

    def on_item_double_clicked(self, item, column):
        item_path = item.data(0, Qt.UserRole)
        if item_path:
            path = Path(item_path)
            if path.is_dir():
                # Expand directory
                if item.childCount() == 0:
                    self.populate_tree_item(item, path)
            elif path.suffix.lower() == '.pdf':
                self.open_pdf_button.setEnabled(True)
                self.process_pdf_button.setEnabled(True)

    def open_selected_pdf(self):
        selected_items = self.file_tree.selectedItems()
        if selected_items:
            item = selected_items[0]
            item_path = item.data(0, Qt.UserRole)
            if item_path:
                path = Path(item_path)
                if path.suffix.lower() == '.pdf':
                    # In a real app, you would open the PDF with system viewer
                    self.statusBar().showMessage(f"Opening: {path.name}")

    def process_selected_pdf(self):
        selected_items = self.file_tree.selectedItems()
        if selected_items:
            item = selected_items[0]
            item_path = item.data(0, Qt.UserRole)
            if item_path:
                path = Path(item_path)
                if path.suffix.lower() == '.pdf':
                    self.file_path = str(path)
                    self.file_label.setText(path.name)
                    self.current_pdf_filename = path.name # Store filename
                    self.tab_widget.setCurrentWidget(self.processor_tab)
                    self.process_pdf()

    def select_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select PDF File", "", "PDF Files (*.pdf)"
        )
        if file_path:
            self.file_path = file_path
            self.current_pdf_filename = os.path.basename(file_path) # Store filename
            self.file_label.setText(os.path.basename(file_path))
            self.process_button.setEnabled(True)
            self.preview_area.clear()
            self.export_button.setEnabled(False)
            self.progress_bar.setValue(0)

    def select_content_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select PDF File", "", "PDF Files (*.pdf)"
        )
        if file_path:
            self.content_file_path = file_path
            self.content_file_label.setText(os.path.basename(file_path))
            self.generate_content_button.setEnabled(True)
            self.content_preview.clear()
            self.export_content_button.setEnabled(False)

    def select_professor_pdf(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select PDF File", "", "PDF Files (*.pdf)"
        )
        if file_path:
            self.professor_file_path = file_path
            self.professor_pdf_label.setText(os.path.basename(file_path))
            self.ask_question_button.setEnabled(True)
            self.answer_area.clear()
            # Extract text for Q&A
            try:
                self.current_pdf_text = self.extract_text_from_pdf(file_path)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to extract text: {str(e)}")

    def extract_text_from_pdf(self, file_path):
        try:
            doc = fitz.open(file_path)
            text = ""
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    def split_text_into_chunks(self, text, max_chunk_size=2500): # Kept for potential use
        """Split text into smaller chunks."""
        text = text.replace('\n', ' ').replace('\r', ' ')
        sentences = [s.strip() for s in text.split('. ') if s.strip()]
        chunks = []
        current_chunk = ""
        for sentence in sentences:
            sentence_with_period = sentence + ". "
            # Check if adding the sentence would exceed the limit
            test_chunk = current_chunk + sentence_with_period
            if len(test_chunk) <= max_chunk_size:
                current_chunk = test_chunk
            else:
                # If current_chunk is not empty, save it
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                # Start a new chunk. If the sentence itself is too long, it will be split.
                if len(sentence_with_period) <= max_chunk_size:
                    current_chunk = sentence_with_period
                else:
                    # Handle very long sentences by forcing a split (less ideal but necessary)
                    # Split the long sentence into parts
                    words = sentence_with_period.split(' ')
                    current_chunk = ""
                    for word in words:
                        if len(current_chunk) + len(word) + 1 <= max_chunk_size:
                             current_chunk += word + " "
                        else:
                            if current_chunk.strip():
                                chunks.append(current_chunk.strip())
                            current_chunk = word + " "
                    # Add the last part of the long sentence
                    if current_chunk.strip():
                        chunks.append(current_chunk.strip())
                    current_chunk = ""
        # Add the final chunk if it exists
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        return chunks

    def process_pdf(self):
        if not hasattr(self, 'file_path'):
            return
        try:
            self.preview_area.append("Extracting text from PDF...")
            text = self.extract_text_from_pdf(self.file_path)
            if not text.strip():
                QMessageBox.warning(self, "Error", "No text found in PDF")
                return
            self.preview_area.append("Processing full text with Kiri...")
            self.process_button.setEnabled(False)
            self.progress_bar.setValue(0)
            # Try to get a topic/title from the first part of the text
            self.current_pdf_topic = self.extract_topic(text[:1000]) # Extract topic from first 1000 chars
            # --- PROCESS ENTIRE TEXT AT ONCE ---
            self.worker = SummarizationWorker(text, "gemma3:4b-it-qat", "summary", is_chunked=False)
            self.worker.output.connect(self.update_preview) # Connect general output
            self.worker.progress.connect(self.progress_bar.setValue)
            self.worker.finished_signal.connect(self.summarization_finished)
            self.worker.start()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to process PDF: {str(e)}")
            self.process_button.setEnabled(True)

    def extract_topic(self, text_snippet):
        """Simple method to extract a potential topic from the beginning of the text."""
        # Try to find a title-like sentence (shorter, possibly capitalized)
        sentences = [s.strip() for s in text_snippet.split('. ') if s.strip()]
        for sentence in sentences[:3]: # Check first 3 sentences
            # Basic heuristic: relatively short and mostly capitalized words
            words = sentence.split()
            if 3 <= len(words) <= 10:
                cap_words = [w for w in words if w.isalpha() and (w.isupper() or w.istitle())]
                if len(cap_words) / len(words) > 0.5: # More than 50% capitalized/Title case
                    return sentence.rstrip('.') # Remove trailing period for title
        # Fallback: first sentence or snippet
        return sentences[0] if sentences else "PDF Content Summary"

    def generate_content(self):
        if not hasattr(self, 'content_file_path'):
            return
        try:
            content_type_map = {
                "Summary": "summary",
                "Brief Overview": "brief",
                "Q&A": "qna",
                "Practice Questions": "questions"
            }
            selected_type = self.content_type_combo.currentText()
            content_type = content_type_map.get(selected_type, "summary")
            # For question types, modify the prompt
            custom_prompt = ""
            if selected_type in ["Q&A", "Practice Questions"]:
                question_count = self.question_count_spin.value()
                question_type = self.question_type_combo.currentText()
                if selected_type == "Q&A":
                    custom_prompt = (
                        f"Generate {question_count} important questions and detailed answers "
                        f"based on the following text. Format each as:\n" # Use \n for newlines
                        f"Q: [Question]\n"
                        f"A: [Detailed answer]\n"
                        f"Question type: {question_type}"
                    )
                else:  # Practice Questions
                    custom_prompt = (
                        f"Create {question_count} practice questions based on the following text. "
                        f"Include {question_type} questions. "
                        f"Number each question clearly."
                    )
            self.content_preview.append(f"Generating {selected_type.lower()}...")
            text = self.extract_text_from_pdf(self.content_file_path)
            if not text.strip():
                QMessageBox.warning(self, "Error", "No text found in PDF")
                return

            # --- PROCESS ENTIRE TEXT AT ONCE for most content types ---
            # Use chunking only if explicitly needed (e.g., for very large texts or specific types)
            # For now, we'll try full text first.
            self.generate_content_button.setEnabled(False)
            self.content_worker = SummarizationWorker(
                text, "gemma3:4b-it-qat", content_type, custom_prompt, is_chunked=False # Try full text
            )
            self.content_worker.output.connect(self.update_content_preview) # Connect general output
            self.content_worker.finished_signal.connect(self.content_generation_finished)
            self.content_worker.start()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate content: {str(e)}")
            self.generate_content_button.setEnabled(True)

    def ask_question(self):
        if not hasattr(self, 'professor_file_path') or not self.current_pdf_text:
            return
        question = self.question_input.toPlainText().strip()
        if not question:
            QMessageBox.warning(self, "Warning", "Please enter a question")
            return
        try:
            self.answer_area.append(f"Q: {question}\n")
            self.answer_area.append("Thinking...\n")
            # Get answer length preference
            length_map = {
                "Brief (1-2 sentences)": "Brief answer (1-2 sentences)",
                "Medium (3-5 sentences)": "Medium-length answer (3-5 sentences)",
                "Detailed (1 paragraph)": "Detailed answer (1 paragraph)",
                "Comprehensive (2+ paragraphs)": "Comprehensive answer (2+ paragraphs)"
            }
            answer_length = length_map.get(self.answer_length_combo.currentText(), "")
            # Create prompt with context (limit context size)
            context_limit = 5000 # Adjust based on model's context window
            prompt = (
                f"Based on the following text, please answer this question: {question}\n"
                f"Context text:\n{self.current_pdf_text[:context_limit]}..."
            )
            # Process in background - Single prompt for Q&A
            self.ask_question_button.setEnabled(False)
            self.qa_worker = SummarizationWorker(
                prompt, "gemma3:4b-it-qat", "professor",
                "You are a helpful professor assistant. Answer questions clearly and provide detailed explanations.",
                answer_length,
                is_chunked=False # Single prompt
            )
            self.qa_worker.output.connect(lambda msg: self.answer_area.append(f"[System] {msg}\n")) # Show system messages
            self.qa_worker.finished_signal.connect(self.qa_finished)
            self.qa_worker.start()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to process question: {str(e)}")
            self.ask_question_button.setEnabled(True)

    def update_preview(self, message):
        """Update the preview area with general processing messages."""
        self.preview_area.append(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {message}")
        self.preview_area.moveCursor(QTextCursor.End) # Auto-scroll to bottom - FIXED

    def update_content_preview(self, message):
        """Update the content preview area with general processing messages."""
        self.content_preview.append(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {message}")
        self.content_preview.moveCursor(QTextCursor.End) # Auto-scroll to bottom - FIXED

    def summarization_finished(self, final_summary):
        self.final_summary = final_summary
        self.process_button.setEnabled(True)
        self.export_button.setEnabled(True)
        self.preview_area.append("\n--- PROCESSING COMPLETE ---\n")
        # Save to notes
        self.save_note(final_summary)
        self.refresh_notes_list()
        QMessageBox.information(self, "Success", "Summary generation completed!")

    def content_generation_finished(self, final_content):
        self.generated_content = final_content
        self.generate_content_button.setEnabled(True)
        self.export_content_button.setEnabled(True)
        self.content_preview.append("\n--- GENERATION COMPLETE ---\n")
        QMessageBox.information(self, "Success", "Content generation completed!")

    def qa_finished(self, answer):
        self.answer_area.append(f"A: {answer}\n")
        self.answer_area.append("---\n")
        self.ask_question_button.setEnabled(True)
        self.question_input.clear()

    def clear_qa(self):
        self.question_input.clear()
        self.answer_area.clear()

    def save_note(self, content):
        note_id = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        note_data = {
            "id": note_id,
            "filename": self.current_pdf_filename if hasattr(self, 'current_pdf_filename') and self.current_pdf_filename else "Unknown",
            "created": datetime.datetime.now().isoformat(),
            "content": content
        }
        self.notes[note_id] = note_data
        self.save_notes()

    def load_notes(self):
        if os.path.exists(self.notes_file):
            try:
                with open(self.notes_file, 'r') as f:
                    data = json.load(f)
                    return data if isinstance(data, dict) else {}
            except Exception:
                return {}
        return {}

    def save_notes(self):
        try:
            with open(self.notes_file, 'w') as f:
                json.dump(self.notes, f, indent=2)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save notes: {str(e)}")

    def refresh_notes_list(self):
        self.notes_list.clear()
        for note_id, note in self.notes.items():
            created = datetime.datetime.fromisoformat(note['created'])
            formatted_date = created.strftime("%Y-%m-%d %H:%M")
            item = QTreeWidgetItem([
                note['filename'],
                formatted_date,
                note_id
            ])
            item.setData(0, Qt.UserRole, note_id)
            self.notes_list.addTopLevelItem(item)

    def on_note_selected(self):
        selected_items = self.notes_list.selectedItems()
        if selected_items:
            item = selected_items[0]
            note_id = item.data(0, Qt.UserRole)
            if note_id in self.notes:
                note = self.notes[note_id]
                self.note_content.setPlainText(note['content'])
                self.export_note_button.setEnabled(True)
                self.delete_note_button.setEnabled(True)
                self.current_note_id = note_id
        else:
            self.note_content.clear()
            self.export_note_button.setEnabled(False)
            self.delete_note_button.setEnabled(False)
            self.current_note_id = None

    def delete_note(self):
        if hasattr(self, 'current_note_id') and self.current_note_id:
            reply = QMessageBox.question(
                self, "Confirm Delete",
                "Are you sure you want to delete this note?",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                del self.notes[self.current_note_id]
                self.save_notes()
                self.refresh_notes_list()
                self.note_content.clear()
                self.export_note_button.setEnabled(False)
                self.delete_note_button.setEnabled(False)
                self.current_note_id = None

    def export_word(self):
        if not hasattr(self, 'final_summary'):
            return
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Summary as Word Document", "", "Word Files (*.docx)"
        )
        if file_path:
            try:
                if not file_path.endswith('.docx'):
                    file_path += '.docx'
                # Pass the topic and filename to the export function
                self.create_word_document(file_path, self.final_summary, self.current_pdf_filename, self.current_pdf_topic)
                QMessageBox.information(self, "Success", "Word document exported successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export Word document: {str(e)}")

    def export_note_word(self):
        if not hasattr(self, 'current_note_id') or not self.current_note_id:
            return
        note = self.notes[self.current_note_id]
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Note as Word Document", "", "Word Files (*.docx)"
        )
        if file_path:
            try:
                if not file_path.endswith('.docx'):
                    file_path += '.docx'
                # For existing notes, we don't have the original topic, so use filename or generic
                note_filename = note.get('filename', 'Unknown')
                note_topic = self.extract_topic(note.get('content', '')[:1000]) if note.get('content') else "Note Content"
                self.create_word_document(file_path, note['content'], note_filename, note_topic)
                QMessageBox.information(self, "Success", "Word document exported successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export Word document: {str(e)}")

    def export_content(self):
        if not hasattr(self, 'generated_content'):
            return
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Content as Word Document", "", "Word Files (*.docx)"
        )
        if file_path:
            try:
                if not file_path.endswith('.docx'):
                    file_path += '.docx'
                content_type = self.content_type_combo.currentText().lower().replace(" ", "_")
                filename = f"generated_{content_type}"
                # For generated content, derive topic from content if possible
                content_topic = self.extract_topic(self.generated_content[:1000]) if self.generated_content else "Generated Content"
                self.create_word_document(file_path, self.generated_content, filename, content_topic)
                QMessageBox.information(self, "Success", "Word document exported successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export Word document: {str(e)}")

    def create_word_document(self, file_path, content, filename="PDF Summary", topic=""):
        """Create a Word document with the content, filename, and topic."""
        doc = Document()
        # Determine document title
        if topic and topic != "PDF Content Summary":
            document_title = topic
        elif filename and filename != "Unknown":
            document_title = filename
        else:
            document_title = "PDF Content Summary"
        # Title
        title = doc.add_heading(document_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        # Metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f'Generated with Gemma3:4b-it-qat\n').bold = True
        if filename and filename != "Unknown":
            metadata.add_run(f'Source: {filename}\n')
        if topic and topic != document_title: # Avoid duplication
             metadata.add_run(f'Topic: {topic}\n')
        metadata.add_run('Editable Notes - Feel free to modify').italic = True
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        # Content
        # Handle potential internal chunking in the result
        if '--- SECTION BREAK ---' in content:
            sections = content.split('--- SECTION BREAK ---')
        else:
            sections = [content]

        for i, section in enumerate(sections):
            if len(sections) > 1:
                section_header = doc.add_heading(f'Section {i+1}', level=1)
                section_header_run = section_header.runs[0]
                section_header_run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
            if "[ERROR]" in section:
                error_para = doc.add_paragraph()
                error_para.add_run(f"[ERROR] {section}").font.color.rgb = RGBColor(244, 67, 54) # Red for errors
            else:
                # For Q&A and Questions, format appropriately
                content_type = "summary" # Default assumption
                if filename and "qna" in filename.lower() or "Q:" in section[:100]:
                    content_type = "qna"
                elif filename and "question" in filename.lower() or any(c.isdigit() and '.' in c for c in section[:200].split('\n')[:5]):
                     content_type = "questions"

                if content_type == "qna":
                     lines = section.split('\n')
                     for line in lines:
                         if line.strip():
                            if line.startswith('Q:'):
                                p = doc.add_paragraph()
                                p.add_run(line).bold = True
                            elif line.startswith('A:'):
                                p = doc.add_paragraph(style='List Bullet')
                                p.add_run(line.lstrip('A: ').strip())
                            else:
                                # Handle potential continuation lines or errors in Q&A format
                                doc.add_paragraph(line)
                elif content_type == "questions":
                    lines = section.split('\n')
                    for line in lines:
                         if line.strip():
                            # Check if it's a numbered list item
                            if line[0].isdigit() and (line[1:3] in ['.', ')'] or (len(line) > 2 and line[2] in ['.', ')'])):
                                p = doc.add_paragraph(style='List Number')
                                # Extract text after number
                                parts = line.split(' ', 1)
                                p.add_run(parts[1] if len(parts) > 1 else line)
                            else:
                                doc.add_paragraph(line)
                else:
                    # Default paragraph/bullet handling for summaries/briefs
                    lines = section.split('\n')
                    for line in lines:
                        if line.strip() and not line.startswith('---'):
                            clean_line = line.strip()
                            if clean_line.startswith('•'):
                                # Main bullet point
                                p = doc.add_paragraph()
                                p.style = doc.styles['List Bullet']
                                p.add_run(clean_line.lstrip('• ')).bold = True
                            elif clean_line.startswith('-'):
                                # Sub-bullet point
                                p = doc.add_paragraph()
                                p.style = doc.styles['List Bullet 2']
                                p.add_run(clean_line.lstrip('- '))
                            elif clean_line:
                                # Regular text
                                doc.add_paragraph(clean_line) # Add as paragraph
            if i < len(sections) - 1:
                doc.add_paragraph()
        doc.save(file_path)

# Added for content generation dialog
from PyQt5.QtWidgets import QInputDialog

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = PDFManagerApp()
    window.show()
    sys.exit(app.exec_())